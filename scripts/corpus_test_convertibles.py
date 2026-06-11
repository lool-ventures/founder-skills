#!/usr/bin/env python3
# /// script
# requires-python = ">=3.10"
# dependencies = ["pdfplumber", "python-docx>=1.2.0"]
# ///
"""Corpus test: Convertible instrument document extraction simulation.

Handles: PDF, DOCX, DOC (via antiword fallback), ZIP (extracts members).

For each document:
  * Loadability + text extraction
  * Document type classification:
    - Israeli CLA (Convertible Loan Agreement) -- typical GKH/Herzog template
    - Israeli Convertible Security / Investment Agreement
    - US Convertible Promissory Note -- NVCA/Cooley/standard forms
    - US Bridge Financing Agreement
  * Jurisdiction detection (Israeli vs Delaware/US)
  * Law firm attribution (current + historical Israeli big-law)
  * Key field extraction:
    - principal / loan amount (USD or NIS)
    - annual interest rate
    - maturity date or months-to-maturity
    - valuation cap / price cap (if any)
    - discount rate (if any)
    - qualified financing threshold
    - conversion mechanics (automatic vs optional)
    - MFN clause
    - interest conversion to shares vs cash repayment

Usage:
    python3 scripts/corpus_test_convertibles.py <CORPUS_DIR> [-o report.json]
"""

from __future__ import annotations

import argparse
import json
import os
import re
import subprocess
import sys
import tempfile
import warnings
import zipfile
from collections import Counter
from pathlib import Path
from typing import Any

# ---- Document type discriminators -------------------------------------------

CLA_MARKERS = [
    r"Convertible Loan Agreement",
    r"\bCLA\b",
    r"convertible loan",
    r"Loan Amount",
]
CONVERTIBLE_NOTE_MARKERS = [
    r"Convertible Promissory Note",
    r"Convertible Note",
    r"Principal Amount",
    r"Promissory Note",
]
CONVERTIBLE_SECURITY_MARKERS = [
    r"Convertible Security",
    r"convertible security",
]
BRIDGE_FINANCING_MARKERS = [
    r"Bridge Financing",
    r"bridge financing",
    r"Bridge Loan",
]
CONVERTIBLE_INVESTMENT_MARKERS = [
    r"Convertible Investment Agreement",
    r"Investment Agreement",
]

# ---- Israeli law firm markers (current + historical) ------------------------
ISRAELI_MARKERS = [
    r"Meitar",
    r"Herzog",
    r"Goldfarb",
    r"Arnon.{0,10}Tadmor",
    r"FISCHER",
    r"Naschitz",
    r"Shibolet",
    r"Amit.{0,10}Pollak",
    r"\bAPM\b",
    r"Erdinast",
    r"\bEBN\b",
    r"Gornitzky",
    r"Barnea",
    r"Pearl Cohen",
    r"H-F\s*&\s*Co",
    r"\bFWMK\b",
    r"\bERM\b",
    r"Firon",
    r"S\.\s*Horowitz",
    # Historical / legacy
    r"GKH",
    r"Gross\s*Kleinhendler",
    r"Yigal Arnon",
    r"Tadmor\s*Levy",
    r"Fischer Behar",
    r"\bFBC\b",
    # Statutory
    r"Israeli law",
    r"Israeli\s+Companies\s+Law",
    r"Tel Aviv",
    r"Section 102",
    r"NIS\b",
    r"New Israeli Shek",
]

DELAWARE_MARKERS = [
    r"State of Delaware",
    r"Delaware corporation",
    r"incorporated in.{0,20}Delaware",
    r"laws of the State of Delaware",
    r"DGCL",
]

# ---- Field extraction patterns ----------------------------------------------
#
# Key corpus findings (from real documents):
# 1. Israeli CLAs: "US$ 7,000,000 (the "Investment Amount")" -- amount BEFORE label
# 2. US notes: "principal amount of US$8,750" -- label BEFORE amount
# 3. US notes: spelled-out "Eight Thousand Dollars (US$8,750)"
# 4. Interest: "at the rate of 8%" (Israeli CLA uses definite article)
# 5. Maturity: in Definitions section as '"Maturity Date" means 12 months'
# 6. Cap: may be called "Price Cap" not "Valuation Cap" in US notes
# 7. Discount: may be "Discounted Price Percentage" = 80% (=> 20% discount)

# Amount before label (Israeli CLA): "US$ 7,000,000 (the 'Investment Amount')"
PRINCIPAL_AMOUNT_FIRST_RE = re.compile(
    r"(?:US\s*)?\$\s*([0-9]{1,3}(?:,[0-9]{3})*(?:\.[0-9]+)?)"
    r"\s*\([^)]{0,40}"
    r"(?:Investment\s+Amount|Principal\s+Amount|Loan\s+Amount|Aggregate\s+Amount)"
    r"[^)]{0,20}\)",
    re.IGNORECASE,
)
# NIS-denominated principal before label
PRINCIPAL_NIS_RE = re.compile(
    r"NIS\s+([0-9]{1,3}(?:,[0-9]{3})*(?:\.[0-9]+)?)"
    r"\s*\([^)]{0,40}"
    r"(?:Investment\s+Amount|Principal\s+Amount|Loan\s+Amount)"
    r"[^)]{0,20}\)",
    re.IGNORECASE,
)
# Label before amount (US notes): "principal amount of US$8,750"
# Also handles spelled-out word amounts before the parenthetical USD value
PRINCIPAL_LABEL_FIRST_RE = re.compile(
    r"(?:principal\s+amount\s+of"
    r"|aggregate\s+principal\s+amount\s+of"
    r"|principal\s+sum\s+of"
    r"|loan\s+amount\s+of"
    r"|investment\s+amount\s+of"
    r"|in\s+the\s+aggregate\s+amount\s+of(?:\s+up\s+to)?)"
    r"\s*(?:[A-Z][a-zA-Z\s,]+\s+)?"
    r"(?:\(?\s*(?:US\s*)?\$\s*([0-9]{1,3}(?:,[0-9]{3})*(?:\.[0-9]+)?)\s*\)?)",
    re.IGNORECASE,
)

# Interest rate -- multiple forms
INTEREST_PATTERNS = [
    # "interest rate ... 8%" / "bears interest at 5%"
    re.compile(
        r"(?:interest\s+rate|bears?\s+interest|accrues?\s+interest)[^\d\n]{0,60}"
        r"([0-9]+(?:\.[0-9]+)?)\s*(?:%|percent)",
        re.IGNORECASE,
    ),
    # "at the rate of 8%" / "at a rate of 5%"  (Israeli CLAs use definite article)
    re.compile(
        r"at\s+(?:the\s+|a\s+)?(?:rate|interest\s+rate)\s+of\s+"
        r"([0-9]+(?:\.[0-9]+)?)\s*(?:%|percent)",
        re.IGNORECASE,
    ),
    # Definitions section: "Interest Rate" means X percent
    # Uses [\x22\x27\u201c\u201d\u2018\u2019] to match straight and smart/curly quotes
    re.compile(
        r"[\x22\x27\u201c\u201d\u2018\u2019]Interest\s+Rate"
        r"[\x22\x27\u201c\u201d\u2018\u2019][^\d\n]{0,40}"
        r"([0-9]+(?:\.[0-9]+)?)\s*(?:%|percent)",
        re.IGNORECASE,
    ),
    # "at a rate of five percent (5%)" -- word-written, numeric in parens
    re.compile(
        r"(?:at\s+(?:the\s+|a\s+)?rate\s+of|bears?\s+interest\s+at)"
        r"\s+\w+(?:\s+\w+)?\s*\(([0-9]+(?:\.[0-9]+)?)\s*%\)",
        re.IGNORECASE,
    ),
]

# Maturity date
MATURITY_DEFN_RE = re.compile(
    # "Maturity Date" means 12 months  (Definitions section)
    r"[\x22\x27\u201c\u201d]Maturity\s+Date[\x22\x27\u201c\u201d]"
    r"\s*[\x22\x27\u201c\u201d]?\s*means\s+"
    r"(?:the\s+date\s+that\s+is\s+)?([0-9]+)\s+months?",
    re.IGNORECASE,
)
MATURITY_DATE_RE = re.compile(
    r"(?:Maturity\s+Date|due\s+and\s+payable\s+(?:on|by))"
    r"[^\d\n]{0,60}"
    r"(?:([A-Z][a-z]+\s+\d{1,2},?\s+\d{4})"
    r"|(\d{1,2}[/-]\d{1,2}[/-]\d{2,4}))",
    re.IGNORECASE,
)
MATURITY_MONTHS_RE = re.compile(
    r"(?:term\s+of\s+|within\s+|period\s+of\s+)"
    r"([0-9]+)\s+months?",
    re.IGNORECASE,
)

# Valuation / Price Cap
CAP_RE = re.compile(
    r"(?:Valuation\s+Cap|Price\s+Cap|Cap\s+Amount|valuation\s+cap)"
    r"[^\$\n]{0,80}(?:US\s*)?\$\s*([0-9]{1,3}(?:,[0-9]{3})*(?:\.[0-9]+)?)",
    re.IGNORECASE,
)
# "$4,000,000 (the 'Price Cap')" -- amount before label
CAP_REVERSED_RE = re.compile(
    r"(?:US\s*)?\$\s*([0-9]{1,3}(?:,[0-9]{3})*(?:\.[0-9]+)?)"
    r"\s*\([^)]{0,30}"
    r"(?:Price\s+Cap|Valuation\s+Cap|Cap\s+Amount)"
    r"[^)]{0,10}\)",
    re.IGNORECASE,
)
# "Company valuation ... equal to US $15,000,000" -- Israeli convertible security form
CAP_VALUATION_RE = re.compile(
    r"company\s+valuation\s+[^$\n]{0,80}(?:US\s*\$|USD\s*|\$)\s*"
    r"([0-9]{1,3}(?:,[0-9]{3})*(?:\.[0-9]+)?)",
    re.IGNORECASE,
)

# Discount rate
DISCOUNT_PATTERNS = [
    re.compile(
        r"[Dd]iscount\s*(?:[Rr]ate)?"
        r"[^\d\n]{0,40}([0-9]+(?:\.[0-9]+)?)\s*(?:%|percent)",
        re.IGNORECASE,
    ),
    # "discount equal to twenty five percent (25%)"
    re.compile(
        r"discount\s+(?:equal\s+to|of)\s+[a-zA-Z\s]+\(([0-9]+(?:\.[0-9]+)?)\s*%\)",
        re.IGNORECASE,
    ),
]
# "Discounted Price Percentage" = 80% => 20% discount (multiplier, not rate)
DISCOUNT_MULTIPLIER_RE = re.compile(
    r"Discounted\s+Price\s+Percentage"
    r"[^\d\n]{0,40}([0-9]+(?:\.[0-9]+)?)\s*(?:%|percent)",
    re.IGNORECASE,
)

# Qualified financing threshold
QUALIFIED_FINANCING_RE = re.compile(
    r"[Qq]ualified\s+[Ff]inancing[^\$\n]{0,80}(?:US\s*)?\$\s*"
    r"([0-9]{1,3}(?:,[0-9]{3})*(?:\.[0-9]+)?)",
    re.IGNORECASE,
)

# Conversion trigger
AUTO_CONVERSION_RE = re.compile(
    r"(?:automatically\s+convert|shall\s+(?:automatically\s+)?convert\s+upon)",
    re.IGNORECASE,
)
OPTIONAL_CONVERSION_RE = re.compile(
    r"(?:may\s+(?:elect\s+to\s+)?convert|at\s+the\s+option\s+of)",
    re.IGNORECASE,
)

# MFN clause
MFN_RE = re.compile(
    r"(?:Most\s+Favored\s+Nation|most\s+favored\s+nation|\bMFN\b)",
    re.IGNORECASE,
)

# Interest converts to shares vs repayment
INTEREST_CONVERTS_RE = re.compile(
    r"(?:accrued\s+interest\s+(?:shall|will)\s+(?:also\s+)?convert"
    r"|interest\s+(?:shall|will)\s+(?:be\s+)?convert(?:ed)?)",
    re.IGNORECASE,
)
INTEREST_REPAY_RE = re.compile(
    r"(?:interest\s+(?:shall|will)\s+(?:be\s+)?(?:repaid|paid\s+in\s+cash)"
    r"|repay(?:ment)?\s+of\s+(?:principal\s+and\s+)?interest)",
    re.IGNORECASE,
)


# ---- Text extraction --------------------------------------------------------


def extract_pdf_text(path: Path) -> str:
    with warnings.catch_warnings():
        warnings.simplefilter("ignore")
        import pdfplumber

        with pdfplumber.open(path) as pdf:
            return "".join((p.extract_text() or "") for p in pdf.pages)


def extract_docx_text(path: Path) -> str:
    import docx

    doc = docx.Document(str(path))
    parts = [p.text for p in doc.paragraphs]
    # Closing binders put content in tables, not paragraphs
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                parts.append(cell.text)
    return "\n".join(parts)


def extract_doc_text(path: Path) -> str:
    """Try antiword / catdoc for legacy .doc files; fall back to empty."""
    for tool in ("antiword", "catdoc"):
        try:
            result = subprocess.run(
                [tool, str(path)],
                capture_output=True,
                text=True,
                timeout=10,
            )
            if result.returncode == 0 and result.stdout.strip():
                return result.stdout
        except (FileNotFoundError, subprocess.TimeoutExpired):
            pass
    # Last-resort: some .doc files are actually DOCX in disguise
    try:
        return extract_docx_text(path)
    except Exception:
        return ""


def extract_zip_members(path: Path) -> list[tuple[str, str]]:
    """Return list of (member_name, text) for PDF/DOCX inside a zip."""
    results = []
    with zipfile.ZipFile(path) as zf:
        for name in zf.namelist():
            suffix = Path(name).suffix.lower()
            if suffix not in {".pdf", ".docx", ".doc"}:
                continue
            try:
                data = zf.read(name)
                fd, tmp_name = tempfile.mkstemp(suffix=suffix)
                tmp = Path(tmp_name)
                with os.fdopen(fd, "wb") as tf:
                    tf.write(data)
                try:
                    if suffix == ".pdf":
                        text = extract_pdf_text(tmp)
                    elif suffix == ".docx":
                        text = extract_docx_text(tmp)
                    else:
                        text = extract_doc_text(tmp)
                    if text.strip():
                        results.append((name, text))
                finally:
                    tmp.unlink(missing_ok=True)
            except Exception:
                pass
    return results


def extract_text(path: Path) -> tuple[str, str]:
    """Return (method, text). method describes how text was extracted."""
    suffix = path.suffix.lower()
    try:
        if suffix == ".pdf":
            return "pdf", extract_pdf_text(path)
        elif suffix in {".docx", ".docm"}:
            return "docx", extract_docx_text(path)
        elif suffix == ".doc":
            text = extract_doc_text(path)
            return ("doc", text) if text.strip() else ("doc_failed", "")
        elif suffix == ".zip":
            members = extract_zip_members(path)
            combined = "\n\n".join(f"=== {n} ===\n{t}" for n, t in members)
            return "zip", combined
        else:
            return "unsupported", ""
    except Exception as e:
        return "failed", f"ERROR: {e}"


# ---- Classification ---------------------------------------------------------


def classify_doc_type(text: str) -> str:
    scores: dict[str, int] = {
        "cla": sum(1 for p in CLA_MARKERS if re.search(p, text, re.IGNORECASE)),
        "convertible_note": sum(1 for p in CONVERTIBLE_NOTE_MARKERS if re.search(p, text, re.IGNORECASE)),
        "convertible_security": sum(1 for p in CONVERTIBLE_SECURITY_MARKERS if re.search(p, text, re.IGNORECASE)),
        "bridge_financing": sum(1 for p in BRIDGE_FINANCING_MARKERS if re.search(p, text, re.IGNORECASE)),
        "convertible_investment": sum(1 for p in CONVERTIBLE_INVESTMENT_MARKERS if re.search(p, text, re.IGNORECASE)),
    }
    best = max(scores, key=lambda k: scores[k])
    return best if scores[best] > 0 else "unknown"


def classify_jurisdiction(text: str) -> str:
    israeli = sum(1 for m in ISRAELI_MARKERS if re.search(m, text, re.IGNORECASE))
    delaware = sum(1 for m in DELAWARE_MARKERS if re.search(m, text, re.IGNORECASE))
    if israeli >= 2 and israeli > delaware:
        return "israeli"
    if delaware >= 2 and delaware > israeli:
        return "delaware"
    if israeli > 0 and delaware == 0:
        return "israeli_likely"
    if delaware > 0 and israeli == 0:
        return "delaware_likely"
    return "unknown"


def detect_law_firm(text: str) -> list[str]:
    firms = {
        "GKH": [r"GKH", r"Gross\s*Kleinhendler"],
        "Herzog": [r"Herzog"],
        "Meitar": [r"Meitar"],
        "Goldfarb": [r"Goldfarb"],
        "Pearl Cohen": [r"Pearl Cohen"],
        "Naschitz": [r"Naschitz"],
        "Arnon/Tadmor": [r"Yigal Arnon", r"Tadmor\s*Levy", r"Arnon.{0,10}Tadmor"],
        "FISCHER/FBC": [r"FISCHER", r"Fischer Behar", r"\bFBC\b"],
        "Shibolet": [r"Shibolet"],
    }
    return [name for name, patterns in firms.items() if any(re.search(p, text, re.IGNORECASE) for p in patterns)]


# ---- Field extraction -------------------------------------------------------


def extract_principal(text: str) -> dict[str, Any] | None:
    # Amount-before-label (Israeli CLA)
    m = PRINCIPAL_AMOUNT_FIRST_RE.search(text)
    if m:
        return {"currency": "USD", "value": float(m.group(1).replace(",", ""))}
    m = PRINCIPAL_NIS_RE.search(text)
    if m:
        return {"currency": "NIS", "value": float(m.group(1).replace(",", ""))}
    # Label-before-amount (US notes)
    m = PRINCIPAL_LABEL_FIRST_RE.search(text)
    if m and m.group(1):
        return {"currency": "USD", "value": float(m.group(1).replace(",", ""))}
    return None


def extract_interest_rate(text: str) -> float | None:
    for pattern in INTEREST_PATTERNS:
        m = pattern.search(text)
        if m:
            try:
                return float(m.group(1))
            except (ValueError, IndexError):
                pass
    return None


def extract_maturity(text: str) -> dict[str, Any] | None:
    m = MATURITY_DEFN_RE.search(text)
    if m:
        return {"type": "months_from_closing", "value": int(m.group(1))}
    m = MATURITY_DATE_RE.search(text)
    if m:
        return {"type": "date", "value": (m.group(1) or m.group(2)).strip()}
    m = MATURITY_MONTHS_RE.search(text)
    if m:
        return {"type": "months_from_closing", "value": int(m.group(1))}
    return None


def extract_cap(text: str) -> float | None:
    for pattern in (CAP_RE, CAP_REVERSED_RE, CAP_VALUATION_RE):
        m = pattern.search(text)
        if m:
            try:
                return float(m.group(1).replace(",", ""))
            except ValueError:
                pass
    return None


def extract_discount(text: str) -> dict[str, Any] | None:
    """Return {rate_pct, multiplier_form}.

    multiplier_form=True means the document stated the multiplier (e.g. 80%)
    and we inverted it to get the discount rate (20%).
    """
    for pattern in DISCOUNT_PATTERNS:
        m = pattern.search(text)
        if m:
            try:
                return {"rate_pct": float(m.group(1)), "multiplier_form": False}
            except ValueError:
                pass
    m = DISCOUNT_MULTIPLIER_RE.search(text)
    if m:
        try:
            return {"rate_pct": round(100 - float(m.group(1)), 4), "multiplier_form": True}
        except ValueError:
            pass
    return None


def extract_qualified_financing(text: str) -> float | None:
    m = QUALIFIED_FINANCING_RE.search(text)
    if m:
        try:
            return float(m.group(1).replace(",", ""))
        except ValueError:
            pass
    return None


def extract_conversion_trigger(text: str) -> str:
    auto = bool(AUTO_CONVERSION_RE.search(text))
    optional = bool(OPTIONAL_CONVERSION_RE.search(text))
    if auto and optional:
        return "both"
    if auto:
        return "automatic"
    if optional:
        return "optional"
    return "not_detected"


def extract_interest_treatment(text: str) -> str:
    converts = bool(INTEREST_CONVERTS_RE.search(text))
    repays = bool(INTEREST_REPAY_RE.search(text))
    if converts and not repays:
        return "converts_to_shares"
    if repays and not converts:
        return "cash_repayment"
    if converts and repays:
        return "both_mentioned"
    return "not_detected"


# ---- Per-file analysis ------------------------------------------------------


def _fields_from_text(text: str) -> dict[str, Any]:
    return {
        "doc_type": classify_doc_type(text),
        "jurisdiction": classify_jurisdiction(text),
        "law_firms": detect_law_firm(text),
        "has_mfn": bool(MFN_RE.search(text)),
        "principal": extract_principal(text),
        "interest_rate_pct": extract_interest_rate(text),
        "maturity": extract_maturity(text),
        "valuation_cap_usd": extract_cap(text),
        "discount": extract_discount(text),
        "qualified_financing_threshold_usd": extract_qualified_financing(text),
        "conversion_trigger": extract_conversion_trigger(text),
        "interest_treatment": extract_interest_treatment(text),
    }


def analyze_file(path: Path) -> dict[str, Any]:
    method, text = extract_text(path)
    text_len = len(text.strip())
    base: dict[str, Any] = {
        "file_size_bytes": path.stat().st_size,
        "extension": path.suffix.lower(),
        "extraction_method": method,
        "text_length": text_len,
    }
    if method in ("failed", "unsupported", "doc_failed") or text_len < 100:
        base["loadable"] = False
        base["error"] = text[:200] if text else "empty/image-only"
        return base
    base["loadable"] = True
    base.update(_fields_from_text(text))
    return base


def analyze_entry(path: Path) -> list[dict[str, Any]]:
    """Return one or more result dicts (zip expands to multiple members)."""
    if path.suffix.lower() == ".zip":
        members = extract_zip_members(path)
        if not members:
            r = analyze_file(path)
            return [r]
        results = []
        for name, text in members:
            r: dict[str, Any] = {
                "file_size_bytes": path.stat().st_size,
                "extension": Path(name).suffix.lower(),
                "extraction_method": "zip_member",
                "text_length": len(text.strip()),
                "loadable": True,
                "zip_member": name,
            }
            r.update(_fields_from_text(text))
            results.append(r)
        return results
    return [analyze_file(path)]


# ---- Main -------------------------------------------------------------------

SUPPORTED_EXTENSIONS = {".pdf", ".docx", ".docm", ".doc", ".zip"}


def main() -> int:
    parser = argparse.ArgumentParser(description=__doc__)
    parser.add_argument("corpus_dir", help="Directory of convertible instrument documents")
    parser.add_argument("-o", "--output", default=None, help="JSON report path")
    parser.add_argument("--verbose", action="store_true")
    args = parser.parse_args()

    corpus = Path(args.corpus_dir)
    if not corpus.is_dir():
        print(f"ERROR: {corpus} is not a directory", file=sys.stderr)
        return 1

    files = sorted(f for f in corpus.iterdir() if f.is_file() and f.suffix.lower() in SUPPORTED_EXTENSIONS)
    print(f"Analyzing {len(files)} documents in {corpus}", file=sys.stderr)

    results: list[dict[str, Any]] = []
    for f in files:
        for entry in analyze_entry(f):
            entry["anon_name"] = f"doc_{len(results):03d}"
            results.append(entry)
            if args.verbose:
                jur = entry.get("jurisdiction", "?")
                dt = entry.get("doc_type", "?")
                firms = ",".join(entry.get("law_firms") or ["?"])
                p = entry.get("principal")
                pr = f"{p['currency']} {p['value']:,.0f}" if p else "--"
                ir = entry.get("interest_rate_pct")
                interest = f"{ir}%" if ir is not None else "--"
                cap = entry.get("valuation_cap_usd")
                cp = f"${cap:,.0f}" if cap else "--"
                disc = entry.get("discount")
                dr = f"{disc['rate_pct']}%" if disc else "--"
                print(
                    f"  [{jur:15s}] {dt:25s} [{firms:20s}]"
                    f"  principal={pr:15s} interest={interest:6s}"
                    f" cap={cp:12s} disc={dr}",
                    file=sys.stderr,
                )

    loadable = [r for r in results if r.get("loadable")]
    n = len(loadable)

    doc_types = Counter(r.get("doc_type", "unknown") for r in loadable)
    jurisdictions = Counter(r.get("jurisdiction", "unknown") for r in loadable)
    law_firms: Counter[str] = Counter()
    for r in loadable:
        for firm in r.get("law_firms") or []:
            law_firms[firm] += 1

    field_hits = {
        "principal": sum(1 for r in loadable if r.get("principal")),
        "interest_rate": sum(1 for r in loadable if r.get("interest_rate_pct") is not None),
        "maturity": sum(1 for r in loadable if r.get("maturity")),
        "valuation_cap": sum(1 for r in loadable if r.get("valuation_cap_usd") is not None),
        "discount": sum(1 for r in loadable if r.get("discount") is not None),
        "qualified_financing": sum(1 for r in loadable if r.get("qualified_financing_threshold_usd") is not None),
    }
    conversion_triggers = Counter(r.get("conversion_trigger", "?") for r in loadable)
    interest_treatments = Counter(r.get("interest_treatment", "?") for r in loadable)

    report = {
        "corpus_dir": str(corpus),
        "total_files": len(files),
        "total_documents": len(results),
        "loadable_count": n,
        "doc_type_distribution": dict(doc_types),
        "jurisdiction_distribution": dict(jurisdictions),
        "law_firm_attribution": dict(law_firms),
        "field_extraction_rates": field_hits,
        "conversion_trigger_distribution": dict(conversion_triggers),
        "interest_treatment_distribution": dict(interest_treatments),
        "mfn_count": sum(1 for r in loadable if r.get("has_mfn")),
        "per_file": results,
    }

    if args.output:
        out_path = os.path.abspath(args.output)
        with open(out_path, "w") as out_f:
            json.dump(report, out_f, indent=2, default=str)
        print(f"\nReport written to {out_path}", file=sys.stderr)

    print("\n" + "=" * 70, file=sys.stderr)
    print("CONVERTIBLE CORPUS TEST SUMMARY", file=sys.stderr)
    print("=" * 70, file=sys.stderr)
    print(
        f"Total files: {len(files)}  |  Documents (incl. zip members): {len(results)}",
        file=sys.stderr,
    )
    print(f"Loadable: {n}/{len(results)}", file=sys.stderr)

    print("\nDocument types:", file=sys.stderr)
    for dt, c in doc_types.most_common():
        print(f"  {c:3d} x {dt}", file=sys.stderr)

    print("\nJurisdiction:", file=sys.stderr)
    for j, c in jurisdictions.most_common():
        print(f"  {c:3d} x {j}", file=sys.stderr)

    print("\nLaw firm attribution:", file=sys.stderr)
    for firm, c in law_firms.most_common():
        print(f"  {c:3d} x {firm}", file=sys.stderr)

    print(f"\nField extraction rates (out of {n} loadable):", file=sys.stderr)
    for field, count in field_hits.items():
        pct = 100 * count / max(n, 1)
        print(f"  {count:3d}/{n} ({pct:5.1f}%) {field}", file=sys.stderr)

    print("\nConversion trigger:", file=sys.stderr)
    for ct, c in conversion_triggers.most_common():
        print(f"  {c:3d} x {ct}", file=sys.stderr)

    print("\nInterest treatment:", file=sys.stderr)
    for it, c in interest_treatments.most_common():
        print(f"  {c:3d} x {it}", file=sys.stderr)

    print(f"\nMFN clause: {report['mfn_count']}/{n}", file=sys.stderr)

    return 0


if __name__ == "__main__":
    sys.exit(main())

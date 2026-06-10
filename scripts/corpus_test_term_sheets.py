#!/usr/bin/env python3
"""Corpus test: Term Sheet / SPA document extraction simulation.

Handles: PDF, DOCX, DOC (via antiword/catdoc fallback).

For each document:
  * Loadability + text extraction
  * Document type: term_sheet (preferred round) | safe | spa | unknown
  * Round series detection (Seed, A, B, C, etc.)
  * Jurisdiction (Israeli vs Delaware)
  * Law firm attribution (current + historical Israeli big-law)
  * Key field extraction:
    - investment_amount (USD)
    - pre_money_valuation (USD)
    - post_money_valuation (USD)
    - price_per_share (USD)
    - option_pool_pct (% target post-financing)
    - liquidation_preference_multiple
    - liquidation_preference_type (participating, non-participating, capped)
    - anti_dilution_type (BBWA, narrow, full_ratchet, none)
    - board_composition (rough: total seats, investor seats)
    - drag_along_threshold (%)
    - tag_along
    - protective_provisions (yes/no)
    - exclusivity_period_days
    - expense_cap_usd
    - investor_name

Usage:
    python3 scripts/corpus_test_term_sheets.py <CORPUS_DIR> [-o report.json]
"""

from __future__ import annotations

import argparse
import json
import re
import subprocess
import sys
import warnings
from collections import Counter
from pathlib import Path
from typing import Any

# ---- Document type discriminators -------------------------------------------

TERM_SHEET_MARKERS = [
    r"Term\s+Sheet",
    r"Summary\s+of\s+Terms",
    r"Non[-\s]?Binding\s+Summary",
    r"Proposed\s+Investment",
    r"Memorandum\s+of\s+Terms",
]
SPA_MARKERS = [
    r"Share\s+Purchase\s+Agreement",
    r"Stock\s+Purchase\s+Agreement",
    r"\bSPA\b",
]
SAFE_MARKERS = [
    # Require the full form name -- "SAFE" alone matches term sheets that mention
    # SAFEs being converted, and SPAs that reference prior SAFEs.
    r"Simple\s+Agreement\s+for\s+Future\s+Equity",
    # Title-line check: SAFE appears in first 500 chars (form title)
    # (handled in classifier code, not regex)
]

# ---- Israeli law firm markers (current + historical) ------------------------
ISRAELI_MARKERS = [
    r"Meitar",
    r"Herzog",
    r"Goldfarb",
    r"Arnon.{0,10}Tadmor",
    r"FISCHER",
    r"Naschitz",
    r"Shibolet",
    r"Amit.{0,10}Pollak",
    r"\bAPM\b",
    r"Erdinast",
    r"\bEBN\b",
    r"Gornitzky",
    r"Barnea",
    r"Pearl Cohen",
    r"H-F\s*&\s*Co",
    r"\bFWMK\b",
    r"\bERM\b",
    r"Firon",
    r"S\.\s*Horowitz",
    # Historical / legacy
    r"GKH",
    r"Gross\s*Kleinhendler",
    r"Yigal Arnon",
    r"Tadmor\s*Levy",
    r"Fischer Behar",
    r"\bFBC\b",
    # Statutory
    r"Israeli law",
    r"Israeli\s+Companies\s+Law",
    r"Tel Aviv",
    r"Section 102",
    r"NIS\b",
    r"New Israeli Shek",
    # Term-sheet specific incorporation phrases
    r"organized\s+under\s+the\s+laws\s+of\s+(?:the\s+State\s+of\s+)?Israel",
    r"incorporated\s+(?:in|under\s+the\s+laws\s+of)\s+(?:the\s+State\s+of\s+)?Israel",
    r"company\s+organized\s+under\s+the\s+laws\s+of\s+Israel",
]

DELAWARE_MARKERS = [
    r"State of Delaware",
    r"Delaware corporation",
    r"incorporated in.{0,20}Delaware",
    r"laws of (?:the State of )?Delaware",
    r"DGCL",
]

# ---- Round series detection -------------------------------------------------
SERIES_PATTERNS = [
    re.compile(r"Series\s+([A-Z]|Seed)(?:\s*-?\s*([0-9]+))?\s+(?:Preferred|Pref)", re.IGNORECASE),
    re.compile(r"Preferred\s+(Seed|[A-Z])(?:\s*-?\s*([0-9]+))?\s+Shares?", re.IGNORECASE),
    re.compile(r"\bSeries\s+([A-Z]|Seed)\s+Financing\b", re.IGNORECASE),
]

# ---- Field extraction patterns ----------------------------------------------
# Money amount helper -- handles "$50 million", "US$25,000,000", "USD 80M", "$1.05B"
# Group 1: numeric value, Group 2: scale word (million|billion|M|B|k)
MONEY_NUM = r"([0-9]{1,4}(?:,[0-9]{3})*(?:\.[0-9]+)?)"
MONEY_SCALE = r"(million|billion|thousand|\bM\b|\bB\b|\bK\b)"
MONEY_RE_INLINE = rf"(?:US\s*\$|USD\s*|\$)\s*{MONEY_NUM}(?:\s*{MONEY_SCALE})?"


def _parse_money(value: str, scale: str | None) -> float:
    n = float(value.replace(",", ""))
    if not scale:
        return n
    s = scale.lower()
    if s in ("billion", "b"):
        return n * 1e9
    if s in ("million", "m"):
        return n * 1e6
    if s in ("thousand", "k"):
        return n * 1e3
    return n


# Investment / round size
INVESTMENT_PATTERNS = [
    # "investment of up to $2,000,000"
    re.compile(
        rf"(?:Investment\s+Amount|investment\s+of(?:\s+up\s+to)?|Amount\s+Raised|Round\s+Size|Total\s+Investment(?:\s+of(?:\s+up\s+to)?)?|total\s+financing\s+round\s+of(?:\s+up\s+to)?)"
        rf"[^\$\n]{{0,40}}{MONEY_RE_INLINE}",
        re.IGNORECASE,
    ),
    # "$25,000,000 (Investment Amount)"
    re.compile(
        rf"{MONEY_RE_INLINE}\s*\([^)]{{0,40}}(?:Investment\s+Amount|Round\s+Size|Total\s+Investment)",
        re.IGNORECASE,
    ),
    # "invest at least $45 million"
    re.compile(
        rf"(?:invest(?:s|ment)?(?:\s+at\s+least)?|will\s+invest)\s+{MONEY_RE_INLINE}",
        re.IGNORECASE,
    ),
]

# Pre-money valuation
PREMONEY_PATTERNS = [
    re.compile(
        rf"(?:fully[-\s]?diluted\s+)?pre[-\s]?money\s+(?:enterprise\s+)?valuation\s+of\s+{MONEY_RE_INLINE}",
        re.IGNORECASE,
    ),
    re.compile(
        rf"pre[-\s]?money[^\$\n]{{0,60}}{MONEY_RE_INLINE}",
        re.IGNORECASE,
    ),
    re.compile(
        rf"{MONEY_RE_INLINE}\s+(?:fully[-\s]?diluted\s+)?pre[-\s]?money",
        re.IGNORECASE,
    ),
]

# Post-money valuation
POSTMONEY_PATTERNS = [
    re.compile(
        rf"(?:fully[-\s]?diluted\s+)?post[-\s]?money\s+(?:enterprise\s+)?valuation\s+of\s+{MONEY_RE_INLINE}",
        re.IGNORECASE,
    ),
    re.compile(
        rf"post[-\s]?money[^\$\n]{{0,60}}{MONEY_RE_INLINE}",
        re.IGNORECASE,
    ),
]

# Option pool percent target
OPTION_POOL_PATTERNS = [
    # "ESOP pool of 7.00% post-money"
    re.compile(
        r"(?:option\s+pool|ESOP(?:\s+pool)?|equity\s+incentive\s+plan|share\s+incentive\s+plan|unallocated\s+(?:share[s]?|pool))"
        r"[^\d\n]{0,80}([0-9]+(?:\.[0-9]+)?)\s*%",
        re.IGNORECASE,
    ),
    # "pool representing 10% of fully diluted"
    re.compile(
        r"pool\s+(?:representing|equal\s+to|of)\s+([0-9]+(?:\.[0-9]+)?)\s*%",
        re.IGNORECASE,
    ),
    # "shall represent 7.00% of the Company's share capital"
    re.compile(
        r"(?:shall\s+represent|increased?\s+to)\s+([0-9]+(?:\.[0-9]+)?)\s*%\s*(?:of\s+(?:the\s+)?(?:Company['']?s\s+)?(?:share\s+capital|fully)|post[-\s])",
        re.IGNORECASE,
    ),
]

# Liquidation preference
LIQ_PREF_MULTIPLE_RE = re.compile(
    r"(?:Liquidation\s+Preference|liquidation\s+preference|preference)"
    r"[^\n]{0,100}?([0-9](?:\.[0-9]+)?)x\s*(?:liquidation\s+)?(?:preference|the\s+(?:Original\s+)?(?:Issue|Purchase)\s+Price|of\s+the\s+(?:Original|original))",
    re.IGNORECASE,
)
# "1x convertible non-participating preference"
LIQ_PREF_INLINE_RE = re.compile(
    r"(?:standard\s+)?([0-9](?:\.[0-9]+)?)x\s+(?:convertible\s+)?(non[-\s]?participating|participating|capped)",
    re.IGNORECASE,
)
LIQ_PARTICIPATION_PATTERNS = {
    "non_participating": re.compile(r"non[-\s]?participating", re.IGNORECASE),
    "participating": re.compile(r"\bfully\s+participating\b|\bparticipating\b", re.IGNORECASE),
    "capped": re.compile(r"participating[^.\n]{0,40}cap(?:ped)?", re.IGNORECASE),
}

# Anti-dilution
AD_PATTERNS = {
    "broad_based_weighted_average": re.compile(
        r"broad[-\s]?based\s+weighted\s+average|broad[-\s]?based\s+anti[-\s]?dilution",
        re.IGNORECASE,
    ),
    "narrow_based_weighted_average": re.compile(r"narrow[-\s]?based\s+weighted\s+average", re.IGNORECASE),
    "full_ratchet": re.compile(r"full\s+ratchet", re.IGNORECASE),
    "weighted_average_unspecified": re.compile(
        r"weighted[-\s]average\s+anti[-\s]?dilution(?!\s+(?:broad|narrow))", re.IGNORECASE
    ),
}

# Board composition
BOARD_PATTERNS = [
    re.compile(
        r"Board\s+(?:of\s+Directors|composition)[^.\n]{0,200}?"
        r"(?:consist\s+of|comprise|shall\s+have|to\s+be\s+composed\s+of)\s+"
        r"([0-9]+|one|two|three|four|five|six|seven)\s+(?:director|member|seat)",
        re.IGNORECASE,
    ),
    # "Insight to designate one director"
    re.compile(
        r"(?:investor[s]?|lead\s+investor|[A-Z][a-zA-Z]+)\s+(?:to\s+designate|shall\s+(?:designate|appoint)|will\s+(?:designate|appoint))\s+"
        r"(one|two|three|[0-9]+)\s+(?:director|board\s+member)",
        re.IGNORECASE,
    ),
]

# Drag-along threshold
DRAG_ALONG_RE = re.compile(
    r"drag[-\s]?along[^.\n]{0,150}?([0-9]+(?:\.[0-9]+)?)\s*%",
    re.IGNORECASE,
)

# Tag-along (presence only)
TAG_ALONG_RE = re.compile(r"tag[-\s]?along|co[-\s]?sale", re.IGNORECASE)

# Protective provisions (presence only)
PROTECTIVE_RE = re.compile(
    r"protective\s+provisions|consent\s+rights|veto\s+rights|matters\s+requiring",
    re.IGNORECASE,
)

# Pro rata / ROFR
PRO_RATA_RE = re.compile(r"pro[-\s]?rata\s+(?:rights?|participation)|preemptive\s+rights", re.IGNORECASE)
ROFR_RE = re.compile(r"right\s+of\s+first\s+refusal|\bROFR\b", re.IGNORECASE)

# Exclusivity / no-shop -- handles "30 days", "thirty (30) days", "Exclusive Negotiations"
EXCLUSIVITY_RE = re.compile(
    r"(?:exclusivity|no[-\s]?shop|exclusive\s+(?:negotiation|dealings))s?"
    r"[^.\n]{0,80}?\(?\s*([0-9]+)\s*\)?[^.\n]{0,5}?(day|week|month)s?",
    re.IGNORECASE,
)

# Expense reimbursement cap
EXPENSE_CAP_RE = re.compile(
    rf"(?:expense[s]?|legal\s+fees|reimburse(?:ment)?)[^\$\n]{{0,80}}{MONEY_RE_INLINE}",
    re.IGNORECASE,
)

# Conversion to common — IPO threshold
IPO_CONVERSION_RE = re.compile(
    rf"(?:underwritten\s+public\s+offering|qualified\s+IPO|QIPO|public\s+offering)"
    rf"[^.\n]{{0,150}}?{MONEY_RE_INLINE}",
    re.IGNORECASE,
)

# ---- Text extraction --------------------------------------------------------


def extract_pdf_text(path: Path) -> str:
    with warnings.catch_warnings():
        warnings.simplefilter("ignore")
        import pdfplumber

        with pdfplumber.open(path) as pdf:
            return "".join((p.extract_text() or "") for p in pdf.pages)


def extract_docx_text(path: Path) -> str:
    import docx

    doc = docx.Document(str(path))
    parts = [p.text for p in doc.paragraphs]
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                parts.append(cell.text)
    return "\n".join(parts)


def extract_doc_text(path: Path) -> str:
    """Legacy .doc -- try antiword / catdoc / textutil; fall back to docx."""
    for tool, args in (
        ("antiword", [str(path)]),
        ("catdoc", [str(path)]),
        ("textutil", ["-convert", "txt", "-stdout", str(path)]),
    ):
        try:
            result = subprocess.run([tool] + args, capture_output=True, text=True, timeout=15)
            if result.returncode == 0 and result.stdout.strip():
                return result.stdout
        except (FileNotFoundError, subprocess.TimeoutExpired):
            pass
    try:
        return extract_docx_text(path)
    except Exception:
        return ""


def extract_text(path: Path) -> str:
    ext = path.suffix.lower()
    if ext == ".pdf":
        return extract_pdf_text(path)
    if ext == ".docx":
        return extract_docx_text(path)
    if ext == ".doc":
        return extract_doc_text(path)
    return ""


# ---- Classification ---------------------------------------------------------


def classify_doc_type(text: str) -> str:
    """Order: SPA > SAFE (explicit) > term_sheet > SAFE-in-title fallback.

    SAFEs and SPAs both reference term sheets, and term sheets often mention
    SAFE conversions. Use the most specific markers first.
    """
    header = text[:500]
    if any(re.search(p, text) for p in SPA_MARKERS):
        return "spa"
    if any(re.search(p, text) for p in SAFE_MARKERS):
        return "safe"
    # SAFE form title heuristic: "SAFE" in header but no other markers
    if re.search(r"\bSAFE\b", header) and "term sheet" not in header.lower():
        return "safe"
    if any(re.search(p, text) for p in TERM_SHEET_MARKERS):
        return "term_sheet"
    return "unknown"


def classify_jurisdiction(text: str) -> str:
    """Returns: israeli | delaware | israeli_likely | delaware_likely | unknown."""
    israeli_hits = sum(1 for p in ISRAELI_MARKERS if re.search(p, text))
    delaware_hits = sum(1 for p in DELAWARE_MARKERS if re.search(p, text))
    if israeli_hits >= 2 and israeli_hits > delaware_hits:
        return "israeli"
    if delaware_hits >= 2 and delaware_hits > israeli_hits:
        return "delaware"
    if israeli_hits >= 1 and delaware_hits == 0:
        return "israeli_likely"
    if delaware_hits >= 1 and israeli_hits == 0:
        return "delaware_likely"
    return "unknown"


def detect_law_firms(text: str) -> list[str]:
    found = []
    for pattern in ISRAELI_MARKERS[:25]:
        m = re.search(pattern, text)
        if m:
            found.append(m.group(0))
    return list(dict.fromkeys(found))[:3]  # dedupe, max 3


def detect_series(text: str) -> str | None:
    for pat in SERIES_PATTERNS:
        m = pat.search(text)
        if m:
            letter = m.group(1)
            sub = m.group(2) if m.lastindex and m.lastindex >= 2 else None
            if sub:
                return f"Series {letter}-{sub}"
            return f"Series {letter}" if letter.lower() != "seed" else "Series Seed"
    return None


# ---- Field extraction -------------------------------------------------------


def _try_money_patterns(patterns: list[re.Pattern], text: str) -> float | None:
    for pat in patterns:
        m = pat.search(text)
        if m:
            try:
                return _parse_money(m.group(1), m.group(2) if m.lastindex and m.lastindex >= 2 else None)
            except (ValueError, IndexError):
                continue
    return None


def extract_investment_amount(text: str) -> float | None:
    return _try_money_patterns(INVESTMENT_PATTERNS, text)


def extract_premoney(text: str) -> float | None:
    return _try_money_patterns(PREMONEY_PATTERNS, text)


def extract_postmoney(text: str) -> float | None:
    return _try_money_patterns(POSTMONEY_PATTERNS, text)


def extract_option_pool(text: str) -> float | None:
    for pat in OPTION_POOL_PATTERNS:
        m = pat.search(text)
        if m:
            try:
                pct = float(m.group(1))
                if 0 < pct < 50:
                    return pct
            except ValueError:
                continue
    return None


def extract_liq_pref(text: str) -> dict[str, Any]:
    multiple = None
    for pat in (LIQ_PREF_INLINE_RE, LIQ_PREF_MULTIPLE_RE):
        m = pat.search(text)
        if m:
            try:
                multiple = float(m.group(1))
                break
            except (ValueError, IndexError):
                continue

    participation_type = None
    # Check capped first (it contains 'participating')
    if LIQ_PARTICIPATION_PATTERNS["capped"].search(text):
        participation_type = "capped"
    elif LIQ_PARTICIPATION_PATTERNS["non_participating"].search(text):
        participation_type = "non_participating"
    elif LIQ_PARTICIPATION_PATTERNS["participating"].search(text):
        participation_type = "participating"

    return {"multiple": multiple, "type": participation_type}


def extract_anti_dilution(text: str) -> str | None:
    for ad_type, pat in AD_PATTERNS.items():
        if pat.search(text):
            return ad_type
    return None


def extract_drag_along(text: str) -> float | None:
    m = DRAG_ALONG_RE.search(text)
    if m:
        try:
            pct = float(m.group(1))
            if 0 < pct <= 100:
                return pct
        except ValueError:
            pass
    return None


def extract_exclusivity_days(text: str) -> int | None:
    m = EXCLUSIVITY_RE.search(text)
    if m:
        try:
            n = int(m.group(1))
            unit = m.group(2).lower()
            if "week" in unit:
                return n * 7
            if "month" in unit:
                return n * 30
            return n
        except ValueError:
            pass
    return None


def extract_expense_cap(text: str) -> float | None:
    # Use the SECOND occurrence if the first is an investment amount
    matches = list(EXPENSE_CAP_RE.finditer(text))
    for m in matches:
        try:
            val = _parse_money(m.group(1), m.group(2) if m.lastindex and m.lastindex >= 2 else None)
            # Expense caps are typically $25k-$250k -- filter out massive investment amounts
            if 5000 <= val <= 500000:
                return val
        except (ValueError, IndexError):
            continue
    return None


def analyze_doc(text: str, label: str) -> dict[str, Any]:
    if not text.strip():
        return {
            "label": label,
            "loadable": False,
            "error": "empty/image-only",
            "doc_type": "unknown",
            "jurisdiction": "unknown",
        }

    doc_type = classify_doc_type(text)
    jurisdiction = classify_jurisdiction(text)
    law_firms = detect_law_firms(text)
    series = detect_series(text)

    investment = extract_investment_amount(text)
    premoney = extract_premoney(text)
    postmoney = extract_postmoney(text)
    option_pool = extract_option_pool(text)
    liq = extract_liq_pref(text)
    ad = extract_anti_dilution(text)
    drag = extract_drag_along(text)
    exclusivity = extract_exclusivity_days(text)
    expense_cap = extract_expense_cap(text)

    return {
        "label": label,
        "loadable": True,
        "text_chars": len(text),
        "doc_type": doc_type,
        "jurisdiction": jurisdiction,
        "law_firms": law_firms,
        "series": series,
        "investment_amount_usd": investment,
        "pre_money_valuation_usd": premoney,
        "post_money_valuation_usd": postmoney,
        "option_pool_pct": option_pool,
        "liq_pref_multiple": liq["multiple"],
        "liq_pref_type": liq["type"],
        "anti_dilution": ad,
        "drag_along_pct": drag,
        "exclusivity_days": exclusivity,
        "expense_cap_usd": expense_cap,
        "has_tag_along": bool(TAG_ALONG_RE.search(text)),
        "has_protective_provisions": bool(PROTECTIVE_RE.search(text)),
        "has_pro_rata": bool(PRO_RATA_RE.search(text)),
        "has_rofr": bool(ROFR_RE.search(text)),
    }


# ---- Report ------------------------------------------------------------------


def _fmt_money(v: float | None) -> str:
    if v is None:
        return "--"
    if v >= 1e9:
        return f"${v / 1e9:.2f}B"
    if v >= 1e6:
        return f"${v / 1e6:.1f}M"
    if v >= 1e3:
        return f"${v / 1e3:.0f}k"
    return f"${v:.0f}"


def main() -> int:
    parser = argparse.ArgumentParser(description=__doc__)
    parser.add_argument("corpus_dir", help="Directory containing term sheet documents")
    parser.add_argument("-o", "--output", help="Write JSON report to file")
    parser.add_argument("--verbose", action="store_true")
    args = parser.parse_args()

    corpus = Path(args.corpus_dir)
    if not corpus.is_dir():
        print(f"Not a directory: {corpus}", file=sys.stderr)
        return 1

    paths = sorted(p for p in corpus.iterdir() if p.suffix.lower() in (".pdf", ".docx", ".doc"))
    print(f"Analyzing {len(paths)} documents in {corpus}")

    results = []
    for p in paths:
        # Anonymize: never store the raw filename in results/report (real
        # company names). Real name goes to stderr only under --verbose.
        label = f"file_{len(results):03d}"
        try:
            text = extract_text(p)
            result = analyze_doc(text, label)
        except Exception as e:
            result = {
                "label": label,
                "loadable": False,
                "error": f"extraction_error: {e}",
                "doc_type": "unknown",
                "jurisdiction": "unknown",
            }
        results.append(result)

        if args.verbose:
            j = result.get("jurisdiction", "?")
            dt = result.get("doc_type", "?")
            s = result.get("series", "?") or "?"
            firms = ",".join(result.get("law_firms") or []) or "?"
            inv = _fmt_money(result.get("investment_amount_usd"))
            pre = _fmt_money(result.get("pre_money_valuation_usd"))
            post = _fmt_money(result.get("post_money_valuation_usd"))
            pool = result.get("option_pool_pct")
            pool_str = f"{pool}%" if pool is not None else "--"
            liq_m = result.get("liq_pref_multiple")
            liq_t = result.get("liq_pref_type")
            liq_str = f"{liq_m}x {liq_t[:8] if liq_t else '?'}" if liq_m is not None else (liq_t or "--")
            print(
                f"  [{j:15}] {dt:11} {s:14} "
                f"inv={inv:8} pre={pre:8} post={post:8} pool={pool_str:6} "
                f"liq={liq_str:18} {firms[:20]} {label}: {p.name[:50]}",
                file=sys.stderr,
            )

    # Summary
    print("\n" + "=" * 70)
    print("TERM-SHEET CORPUS TEST SUMMARY")
    print("=" * 70)
    total = len(results)
    loadable = [r for r in results if r.get("loadable")]
    print(f"Total files: {total}  |  Loadable: {len(loadable)}/{total}")

    print("\nDocument types:")
    for dt, n in Counter(r["doc_type"] for r in loadable).most_common():
        print(f"  {n:3} x {dt}")

    print("\nJurisdiction:")
    for j, n in Counter(r["jurisdiction"] for r in loadable).most_common():
        print(f"  {n:3} x {j}")

    print("\nLaw firm attribution:")
    firms_counter: Counter[str] = Counter()
    for r in loadable:
        for f in r.get("law_firms") or []:
            firms_counter[f] += 1
    for f, n in firms_counter.most_common():
        print(f"  {n:3} x {f}")

    print("\nRound series:")
    for s, n in Counter(r.get("series") or "unknown" for r in loadable).most_common():
        print(f"  {n:3} x {s}")

    if loadable:
        print(f"\nField extraction rates (out of {len(loadable)} loadable):")
        fields = [
            ("investment_amount_usd", "investment_amount"),
            ("pre_money_valuation_usd", "pre_money"),
            ("post_money_valuation_usd", "post_money"),
            ("option_pool_pct", "option_pool_pct"),
            ("liq_pref_multiple", "liq_pref_multiple"),
            ("liq_pref_type", "liq_pref_type"),
            ("anti_dilution", "anti_dilution"),
            ("drag_along_pct", "drag_along_pct"),
            ("exclusivity_days", "exclusivity_days"),
            ("expense_cap_usd", "expense_cap_usd"),
        ]
        for key, label in fields:
            n = sum(1 for r in loadable if r.get(key) is not None)
            pct = 100 * n / len(loadable)
            print(f"  {n:2}/{len(loadable)} ({pct:5.1f}%) {label}")

        # Boolean flags
        print("\nBoolean signals (out of loadable):")
        flags = [
            ("has_tag_along", "tag_along"),
            ("has_protective_provisions", "protective_provisions"),
            ("has_pro_rata", "pro_rata"),
            ("has_rofr", "ROFR"),
        ]
        for key, label in flags:
            n = sum(1 for r in loadable if r.get(key))
            print(f"  {n:2}/{len(loadable)} {label}")

        # Anti-dilution distribution
        print("\nAnti-dilution distribution:")
        for ad, n in Counter(r.get("anti_dilution") or "not_detected" for r in loadable).most_common():
            print(f"  {n:3} x {ad}")

    if args.output:
        Path(args.output).write_text(json.dumps(results, indent=2, default=str))
        print(f"\nReport written to {args.output}")
    return 0


if __name__ == "__main__":
    sys.exit(main())
